"""Build deterministic Phase 2 PDF, DOCX, and HTML acceptance fixtures."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "tests" / "format_fixtures"


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def build_docx(path: Path, *, mandatory: bool, days: int) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    colors = {1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}
    sizes = {1: 16, 2: 13, 3: 12}
    spacing = {1: (18, 10), 2: (14, 7), 3: (10, 5)}
    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(sizes[level])
        style.font.color.rgb = RGBColor.from_string(colors[level])
        style.paragraph_format.space_before = Pt(spacing[level][0])
        style.paragraph_format.space_after = Pt(spacing[level][1])

    header = section.header.paragraphs[0]
    header.text = "GovernDiff Sample Policy - running header"
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    _add_page_field(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("GovernDiff Digital Policy Fixture")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor.from_string("1F3A5F")
    doc.add_paragraph("A deterministic DOCX used to verify headings, lists, tables, and evidence indices.")

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph("This policy applies to all project members.")
    doc.add_paragraph("Submit reports through the documented review channel.", style="List Bullet")
    doc.add_paragraph("Keep evidence with every decision.", style="List Bullet")

    doc.add_heading("2. Duties", level=1)
    modal = "must" if mandatory else "may"
    doc.add_paragraph(f"The Board {modal} respond within {days} days.")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    values = [
        ("Role", "Decision", "Deadline"),
        ("Board", "Review exception", f"{days} days"),
        ("Secretary", "Publish notice", "5 days"),
    ]
    for row, values_row in zip(table.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    _set_table_geometry(table, [1800, 4500, 3060])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("3. Records", level=1)
    doc.add_paragraph("Records must retain the source paragraph and review outcome.")
    doc.save(path)


def build_pdf(path: Path, *, mandatory: bool, days: int) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter, pageCompression=0)
    width, height = letter
    pages = [
        ["Table of Contents", "1. Scope ........................ 2", "2. Duties ....................... 3"],
        ["1. Scope", "This policy applies to all project members.", "- Submit reports through the documented review channel."],
        ["2. Duties", f"The Board {'must' if mandatory else 'may'} respond within {days} days."],
    ]
    for page_number, lines in enumerate(pages, start=1):
        pdf.setFont("Helvetica", 9)
        pdf.drawString(72, height - 36, "GovernDiff Sample Policy - running header")
        pdf.drawRightString(width - 72, 36, f"Page {page_number} of {len(pages)}")
        y = height - 90
        for index, line in enumerate(lines):
            pdf.setFont("Helvetica-Bold" if index == 0 else "Helvetica", 14 if index == 0 else 11)
            pdf.drawString(72, y, line)
            y -= 30
        pdf.showPage()
    pdf.save()


def build_html(path: Path, *, mandatory: bool, days: int) -> None:
    modal = "must" if mandatory else "may"
    value = f"""<!doctype html>
<html lang="en">
<head><title>GovernDiff fixture</title><style>.hidden {{display:none}}</style><script>window.noise = true;</script></head>
<body>
  <header>Site masthead that must be ignored</header>
  <nav><a href="/">Home</a><a href="/news">News</a></nav>
  <main>
    <div id="table-of-contents"><h2>Contents</h2><ol><li>Scope</li><li>Duties</li></ol></div>
    <h1>1. Scope</h1>
    <p>This policy applies to all project members.</p>
    <ul><li>Submit reports through the documented review channel.</li><li>Keep evidence with every decision.</li></ul>
    <h1>2. Duties</h1>
    <p>The Board {modal} respond within {days} days.</p>
    <table>
      <thead><tr><th>Role</th><th>Decision</th><th>Deadline</th></tr></thead>
      <tbody><tr><td>Board</td><td>Review exception</td><td>{days} days</td></tr><tr><td>Secretary</td><td>Publish notice</td><td>5 days</td></tr></tbody>
    </table>
  </main>
  <footer>Footer navigation that must be ignored</footer>
</body></html>
"""
    path.write_text(value, encoding="utf-8", newline="\n")


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_pdf(OUT / "digital_policy_old.pdf", mandatory=False, days=30)
    build_pdf(OUT / "digital_policy_new.pdf", mandatory=True, days=10)
    build_docx(OUT / "policy_old.docx", mandatory=False, days=30)
    build_docx(OUT / "policy_new.docx", mandatory=True, days=10)
    build_html(OUT / "policy_old.html", mandatory=False, days=30)
    build_html(OUT / "policy_new.html", mandatory=True, days=10)
    (OUT / "toc_noise_old.html").write_text(
        "<main><div class='toc'><h2>Contents</h2><ol><li>Scope</li></ol></div>"
        "<h1>Policy</h1><p>Members must keep records.</p></main>",
        encoding="utf-8",
        newline="\n",
    )
    (OUT / "toc_noise_new.html").write_text(
        "<main><div class='toc'><h2>Contents</h2><ol><li>Scope and duties</li></ol></div>"
        "<h1>Policy</h1><p>Members must keep records.</p></main>",
        encoding="utf-8",
        newline="\n",
    )

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with (OUT / "scanned_like.pdf").open("wb") as handle:
        writer.write(handle)

    mixed = PdfWriter()
    mixed.add_page(PdfReader(OUT / "digital_policy_old.pdf").pages[1])
    mixed.add_blank_page(width=612, height=792)
    mixed.add_blank_page(width=612, height=792)
    with (OUT / "mixed_text_coverage.pdf").open("wb") as handle:
        mixed.write(handle)

    encrypted = PdfWriter()
    for page in PdfReader(OUT / "digital_policy_old.pdf").pages:
        encrypted.add_page(page)
    encrypted.encrypt("governdiff-test")
    with (OUT / "encrypted.pdf").open("wb") as handle:
        encrypted.write(handle)

    (OUT / "corrupt.pdf").write_bytes(b"%PDF-1.7\nnot a complete PDF")
    (OUT / "corrupt.docx").write_bytes(b"not a zip-based Office document")
    (OUT / "empty.txt").write_bytes(b"")
    (OUT / "garbled.txt").write_text("\ufffd\ufffd\ufffd\ufffd readable \ufffd\ufffd\ufffd\ufffd", encoding="utf-8")


if __name__ == "__main__":
    build()
    print(OUT)
